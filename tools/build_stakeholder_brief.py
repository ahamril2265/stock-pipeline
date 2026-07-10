from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs")
OUT_PATH = OUT_DIR / "stock_pipeline_stakeholder_brief.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph_border_bottom(paragraph, color="D9E2F3", size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_para(doc, text="", style=None, size=11, color=BLACK, bold=False, italic=False, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color="D9E2F3", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=INK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=BLACK)
    add_para(doc, "", after=4)


def add_label_detail_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660], 120)
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        c1, c2 = table.rows[idx].cells
        set_cell_shading(c1, LIGHT_GRAY)
        for cell in (c1, c2):
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        r1 = c1.paragraphs[0].add_run(label)
        set_run_font(r1, size=10, color=INK, bold=True)
        r2 = c2.paragraphs[0].add_run(value)
        set_run_font(r2, size=10, color=BLACK)
    add_para(doc, "", after=4)
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    set_table_geometry(table, widths, 120)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_GRAY)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=BLACK)
    set_table_geometry(table, widths, 120)
    add_para(doc, "", after=4)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    r = footer_p.add_run("Stock Analytics Pipeline stakeholder brief")
    set_run_font(r, size=9, color=MUTED)


def add_opening(doc):
    add_para(doc, "STAKEHOLDER BRIEF", size=10, color=MUTED, bold=True, after=2)
    title = add_para(
        doc,
        "Real-Time Stock Market Analytics Pipeline",
        size=24,
        color=INK,
        bold=True,
        after=4,
    )
    title.paragraph_format.space_before = Pt(14)
    add_para(
        doc,
        "A containerized streaming lakehouse platform for ingesting market events, "
        "producing governed analytics layers, and serving dashboard-ready insights.",
        size=13,
        color=MUTED,
        after=14,
    )
    add_label_detail_table(
        doc,
        [
            ("Audience", "Business, product, analytics, and engineering stakeholders"),
            ("Prepared by", "Ahamed Rilwan Mohaaideen"),
            ("Date", date.today().strftime("%B %d, %Y")),
            ("Status", "Implemented local end-to-end data platform with dashboard and orchestration components"),
        ],
    )
    rule = add_para(doc, "", after=8)
    paragraph_border_bottom(rule, color="D9E2F3", size="8", space="4")


def build_document():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_opening(doc)

    add_heading(doc, "Executive Snapshot", 1)
    add_callout(
        doc,
        "Stakeholder takeaway",
        "This project demonstrates how a real-time stock analytics platform can move from raw market events "
        "to trusted, dashboard-ready metrics using a modern data engineering architecture.",
    )
    add_para(
        doc,
        "The system simulates market activity, publishes price ticks and trade executions into Kafka, "
        "governs the payloads through Avro schemas, processes the data with Spark, stores curated Delta Lake "
        "datasets in MinIO, and serves Gold-layer analytics through ClickHouse and a Streamlit dashboard.",
    )
    for item in [
        "Shows end-to-end data movement from streaming ingestion to executive analytics.",
        "Uses the Bronze, Silver, and Gold pattern to separate raw retention, validated metrics, and serving tables.",
        "Includes operational concerns that matter to production stakeholders: schema management, validation, quarantine paths, replay, optimization, orchestration, and monitoring views.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Project At A Glance", 1)
    add_label_detail_table(
        doc,
        [
            ("Problem addressed", "Stakeholders need timely, trusted visibility into market activity and pipeline health."),
            ("Primary users", "Analytics teams, data engineers, product stakeholders, and reviewers evaluating platform readiness."),
            ("Core outcome", "Dashboard-ready market KPIs, symbol summaries, volume rankings, OHLC data, and operational status views."),
            ("Current deployment", "Docker Compose based local environment with Kafka, Schema Registry, Spark, MinIO, ClickHouse, Airflow, Postgres, and Streamlit."),
            ("Presentation angle", "A practical proof of modern streaming architecture and governed analytical data products."),
        ],
    )

    add_heading(doc, "Architecture Narrative", 1)
    add_para(
        doc,
        "The architecture is organized as a governed streaming lakehouse. Each stage has a clear responsibility, "
        "which makes the system easier to explain, operate, and extend.",
    )
    add_matrix_table(
        doc,
        ["Stage", "Role", "Stakeholder value"],
        [
            ("Producer", "Generates price tick and trade execution events for AAPL, TSLA, NVDA, MSFT, and AMZN.", "Provides a repeatable real-time demo without relying on external market feeds."),
            ("Kafka + Schema Registry", "Receives Avro messages and manages schemas for price and trade event contracts.", "Creates a governed event interface between producers and consumers."),
            ("Bronze Delta Lake", "Decodes Kafka records and stores raw events in MinIO by event date and namespace.", "Preserves replayable source data and supports auditability."),
            ("Silver Delta Lake", "Validates, deduplicates, timestamps, and aggregates events into business metrics.", "Turns raw events into trusted analytical building blocks."),
            ("Gold + ClickHouse", "Publishes curated tables for symbol summary, market KPIs, top symbols, and OHLC.", "Serves fast dashboard and BI-style queries."),
            ("Dashboard + Airflow", "Exposes market, symbol, health, storage, Spark, Kafka, Airflow, and benchmark views while orchestrating batch refreshes hourly.", "Gives stakeholders both business insight and operational visibility."),
        ],
        [1500, 4250, 3610],
    )

    add_heading(doc, "Business And Analytical Outputs", 1)
    add_para(
        doc,
        "The Gold layer is the main stakeholder-facing product surface. It translates streaming activity into concise "
        "tables that a dashboard, analyst workflow, or BI layer can consume.",
    )
    add_matrix_table(
        doc,
        ["Output", "Metrics included", "How stakeholders use it"],
        [
            ("gold_market_kpis", "Total market volume, buy/sell volume, market VWAP, average market price, average latency, active symbols.", "Executive summary of market activity and system freshness."),
            ("gold_symbol_summary", "Latest price, daily volume, VWAP, average spread, average latency, buy volume, sell volume.", "Per-symbol performance review and comparison."),
            ("gold_top_symbols", "Volume ranking, VWAP, latest price, buy/sell volume.", "Leaderboard for the most actively traded symbols."),
            ("gold_ohlc", "Open, high, low, close, and volume.", "Candlestick visualization and trend analysis."),
        ],
        [2100, 4300, 2960],
    )

    add_heading(doc, "Data Governance And Reliability", 1)
    add_para(
        doc,
        "Several design choices make the project stronger than a simple demo pipeline:",
    )
    for item in [
        "Schema governance: price and trade events are encoded as Avro records and registered with Confluent Schema Registry.",
        "Validation controls: Silver processors filter invalid prices, volumes, spreads, and latency values before creating analytical datasets.",
        "Quarantine handling: failed decode and invalid business records are written to quarantine locations instead of silently disappearing.",
        "Replay and maintenance: the repo includes a Silver replay workflow and Delta optimization script for recovery and storage hygiene.",
        "Workflow orchestration: Airflow runs the Silver and Gold builders on an hourly schedule with explicit dependencies.",
        "Graceful operation: producer and streaming jobs include shutdown handling to reduce abrupt process failure during demos.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Dashboard Story For Stakeholders", 1)
    add_para(
        doc,
        "A presentation can be framed as a short journey through business value and platform trust.",
    )
    for step in [
        "Start with Market Overview to show business KPIs and the live analytical surface.",
        "Move to Top Symbols and Symbol Analysis to demonstrate drill-down from market-wide signal to per-symbol insight.",
        "Open OHLC to show chart-ready Gold data for price movement analysis.",
        "Switch to Pipeline Health, Spark, Kafka, Airflow, and Storage views to show operational observability.",
        "Close with Performance Benchmark as a readiness discussion, while noting that some benchmark values are demo-oriented and generated in the dashboard layer.",
    ]:
        add_number(doc, step)

    add_heading(doc, "Current Capability Assessment", 1)
    add_matrix_table(
        doc,
        ["Area", "Implemented capability", "Stakeholder message"],
        [
            ("Ingestion", "Unified Avro producer creates price ticks and trade events and publishes to Kafka.", "The system can simulate real-time market flow end to end."),
            ("Processing", "Spark handles Bronze decoding and Silver/Gold transformations.", "The pipeline separates raw capture from business logic and serving outputs."),
            ("Storage", "Delta Lake on MinIO stores Bronze and Silver datasets.", "Data remains replayable and organized for downstream rebuilds."),
            ("Serving", "ClickHouse hosts Gold tables consumed by the dashboard.", "Curated analytics are optimized for interactive access."),
            ("Operations", "Airflow DAG, health checks, logs, and monitoring pages are present.", "The project speaks to runability, not just data transformation."),
        ],
        [1750, 4300, 3310],
    )

    add_heading(doc, "Risks, Constraints, And Discussion Points", 1)
    for item in [
        "The current environment is a local Docker Compose deployment, so production deployment would require security hardening, secrets management, scaling design, and infrastructure automation.",
        "The market feed is simulated, which is appropriate for demonstration but should be replaced with a licensed or approved real market data source for production use.",
        "Dashboard benchmark metrics include generated values in the application layer, so they should be treated as presentation indicators until connected to measured telemetry.",
        "Some service credentials are suitable for local development only and should be moved to environment-specific secret management before wider deployment.",
        "Automated tests, CI/CD, and data quality contracts would strengthen confidence before stakeholder sign-off for a production roadmap.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Recommended Next Steps", 1)
    add_matrix_table(
        doc,
        ["Priority", "Action", "Expected impact"],
        [
            ("1", "Connect dashboard performance and health metrics to measured service telemetry.", "Improves trust in the monitoring layer."),
            ("2", "Add CI checks for schema registration, Spark transformations, and ClickHouse loaders.", "Reduces regression risk as the project evolves."),
            ("3", "Externalize credentials and configuration for non-local environments.", "Prepares the system for safer stakeholder demos and deployment reviews."),
            ("4", "Integrate a real or recorded market data source.", "Moves the platform closer to production realism."),
            ("5", "Package a guided demo script with expected service URLs and talking points.", "Makes stakeholder presentations repeatable and crisp."),
        ],
        [950, 4750, 3660],
    )

    add_heading(doc, "Presentation Close", 1)
    add_callout(
        doc,
        "Recommended closing message",
        "This project is a practical demonstration of a governed, observable, real-time data platform. "
        "It shows how streaming events can become trusted analytics products through clear architectural layers "
        "and operational controls.",
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path.resolve())
